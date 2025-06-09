import os
from typing import List, Optional, Tuple, BinaryIO, Dict, Any
from fastapi import UploadFile, HTTPException
from sqlalchemy.orm import Session
from fastapi.responses import StreamingResponse
import mimetypes
import uuid
from datetime import datetime
import shutil
from pathlib import Path
import asyncio
from docx import Document
import pandas as pd
import openpyxl

from app.models.core.user import User
from app.models.reports import (
    Report,
    ReportType,
    ReportStatus,
    ReportTypeCategory,
    AnalysisType,
    MetadataType,
    ReportInsight,
    ReportQuery,
    ReportContent,
    ReportVersion
)
from app.schemas.report import (
    ReportCreate,
    ReportUpdate,
    ReportResponse,
    ReportTypeResponse,
    ReportStatusResponse
)
from app.repositories.report import report_repository
from app.config.settings import get_settings
from app.core.exceptions import (
    NotFoundException,
    PermissionException,
    ValidationException,
    DatabaseError,
    AIProcessingError
)

settings = get_settings()


class ReportService:
    """Service for handling report operations."""

    def __init__(self, db: Session):
        self.db = db
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.max_upload_size = settings.MAX_UPLOAD_SIZE
        self.allowed_extensions = [ext.strip() for ext in settings.ALLOWED_EXTENSIONS.split(",")]
        self.cache_dir = Path(settings.CACHE_DIR)
        self.cache_ttl = settings.CACHE_TTL
        
        # Ensure directories exist
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        self.cache_dir.mkdir(parents=True, exist_ok=True)

    async def create_report(self, db: Session, report_in: ReportCreate, file: UploadFile, user: User) -> ReportResponse:
        """Create a new report."""
        try:
            report_data = report_in.dict()
            report_data["created_by"] = user.id
            report_data["updated_by"] = user.id
            
            # Validate file
            if not file:
                raise ValidationException("No file provided")
            
            # Check file size
            if file.size > self.max_upload_size:
                raise ValidationException(f"File too large. Maximum size is {self.max_upload_size} bytes")
            
            # Check file extension
            file_ext = os.path.splitext(file.filename)[1].lower().lstrip('.')
            if file_ext not in self.allowed_extensions:
                raise ValidationException(f"File type not allowed. Allowed types: {', '.join(self.allowed_extensions)}")
            
            # Create report
            report = Report(
                title=report_in.title,
                description=report_in.description,
                user_id=user.id,
                file_path=str(self.upload_dir / f"{uuid.uuid4()}_{file.filename}"),
                file_type=file_ext,
                file_size=file.size,
                status="pending"
            )
            
            # Save file
            try:
                with open(report.file_path, "wb") as f:
                    await file.seek(0)
                    while chunk := await file.read(1024 * 1024):  # 1MB chunks
                        f.write(chunk)
            except IOError as e:
                raise ValidationException(f"Error saving file: {str(e)}")
            
            # Save report to database
            try:
                report = report_repository.create(self.db, obj_in=report)
            except DatabaseError as e:
                # Clean up file if database operation fails
                if os.path.exists(report.file_path):
                    os.remove(report.file_path)
                raise DatabaseError(f"Error saving report to database: {str(e)}")
            
            # Process report asynchronously
            try:
                asyncio.create_task(self.process_report(self.db, report.id))
            except Exception as e:
                logger.error(f"Error scheduling report processing: {str(e)}")
                # Don't raise here as the report is already created
            
            return ReportResponse.from_orm(report)
        except ValidationException as e:
            raise e
        except DatabaseError as e:
            raise e
        except Exception as e:
            logger.error(f"Unexpected error creating report: {str(e)}", exc_info=True)
            raise DatabaseError(f"Error creating report: {str(e)}")

    async def list_reports(
        self,
        user: User,
        skip: int = 0,
        limit: int = 100,
        type: Optional[ReportType] = None,
        category: Optional[ReportTypeCategory] = None,
        status: Optional[ReportStatus] = None,
        analysis_type: Optional[AnalysisType] = None,
        metadata_type: Optional[MetadataType] = None,
        is_archived: Optional[bool] = None,
        is_public: Optional[bool] = None
    ) -> List[ReportResponse]:
        """List reports with optional filters."""
        query = self.db.query(Report)
        
        # Apply filters
        if type:
            query = query.filter(Report.type == type)
        if category:
            query = query.filter(Report.category == category)
        if status:
            query = query.filter(Report.status == status)
        if analysis_type:
            query = query.filter(Report.analysis_type == analysis_type)
        if metadata_type:
            query = query.filter(Report.metadata_type == metadata_type)
        if is_archived is not None:
            query = query.filter(Report.is_archived == is_archived)
        if is_public is not None:
            query = query.filter(Report.is_public == is_public)
            
        # Get user's reports and shared reports
        reports = (
            query.filter(
                (Report.created_by == user.id) |
                (Report.is_public == True)
            )
            .offset(skip)
            .limit(limit)
            .all()
        )
        return [ReportResponse.from_orm(report) for report in reports]

    async def get_report(
        self,
        user: User,
        report_id: uuid.UUID
    ) -> Optional[ReportResponse]:
        """Get a specific report."""
        report = report_repository.get(self.db, id=report_id)
        if not report:
            return None
            
        # Check access
        if report.created_by != user.id and not report.is_public:
            return None
            
        return ReportResponse.from_orm(report)

    async def update_report(
        self,
        user: User,
        report_id: uuid.UUID,
        report_in: ReportUpdate,
        file: Optional[UploadFile] = None
    ) -> Optional[ReportResponse]:
        """Update a report."""
        report = report_repository.get(self.db, id=report_id)
        if not report:
            return None
            
        # Check access
        if report.created_by != user.id:
            return None

        update_data = report_in.dict(exclude_unset=True)
        update_data["updated_by"] = user.id
        
        # Handle file upload if provided
        if file:
            await self._handle_file_upload(report_id, file)
        
        report = report_repository.update(self.db, db_obj=report, obj_in=update_data)
        return ReportResponse.from_orm(report)

    async def delete_report(
        self,
        user: User,
        report_id: uuid.UUID
    ) -> bool:
        """Delete a report."""
        report = report_repository.get(self.db, id=report_id)
        if not report:
            return False
            
        # Check access
        if report.created_by != user.id:
            return False

        # Delete associated files
        await self._delete_report_files(report_id)
        
        report_repository.remove(self.db, id=report_id)
        return True

    async def archive_report(
        self,
        user: User,
        report_id: uuid.UUID
    ) -> Optional[ReportResponse]:
        """Archive a report."""
        report = report_repository.get(self.db, id=report_id)
        if not report:
            return None
            
        # Check access
        if report.created_by != user.id:
            return None

        update_data = {"is_archived": True, "updated_by": user.id}
        report = report_repository.update(self.db, db_obj=report, obj_in=update_data)
        return ReportResponse.from_orm(report)

    async def unarchive_report(
        self,
        user: User,
        report_id: uuid.UUID
    ) -> Optional[ReportResponse]:
        """Unarchive a report."""
        report = report_repository.get(self.db, id=report_id)
        if not report:
            return None
            
        # Check access
        if report.created_by != user.id:
            return None

        update_data = {"is_archived": False, "updated_by": user.id}
        report = report_repository.update(self.db, db_obj=report, obj_in=update_data)
        return ReportResponse.from_orm(report)

    async def make_public(
        self,
        user: User,
        report_id: uuid.UUID
    ) -> Optional[ReportResponse]:
        """Make a report public."""
        report = report_repository.get(self.db, id=report_id)
        if not report:
            return None
            
        # Check access
        if report.created_by != user.id:
            return None

        update_data = {"is_public": True, "updated_by": user.id}
        report = report_repository.update(self.db, db_obj=report, obj_in=update_data)
        return ReportResponse.from_orm(report)

    async def make_private(
        self,
        user: User,
        report_id: uuid.UUID
    ) -> Optional[ReportResponse]:
        """Make a report private."""
        report = report_repository.get(self.db, id=report_id)
        if not report:
            return None
            
        # Check access
        if report.created_by != user.id:
            return None

        update_data = {"is_public": False, "updated_by": user.id}
        report = report_repository.update(self.db, db_obj=report, obj_in=update_data)
        return ReportResponse.from_orm(report)

    async def list_report_types(self) -> List[ReportTypeResponse]:
        """List all report types."""
        report_types = self.db.query(ReportType).all()
        return [ReportTypeResponse.from_orm(rt) for rt in report_types]

    async def list_report_statuses(self) -> List[ReportStatusResponse]:
        """List all report statuses."""
        report_statuses = self.db.query(ReportStatus).all()
        return [ReportStatusResponse.from_orm(rs) for rs in report_statuses]

    async def get_file_content(
        self,
        user: User,
        report_id: int
    ) -> Tuple[BinaryIO, str, str]:
        """Get file content for streaming.
        
        Returns:
            Tuple[BinaryIO, str, str]: File object, filename, and content type
        """
        report = (
            self.db.query(Report)
            .filter(Report.id == report_id, Report.user_id == user.id)
            .first()
        )
        if not report:
            raise HTTPException(status_code=404, detail="Report not found")

        if not os.path.exists(report.file_path):
            raise HTTPException(status_code=404, detail="File not found")

        # Get content type
        content_type, _ = mimetypes.guess_type(report.file_path)
        if not content_type:
            content_type = "application/octet-stream"

        # Get original filename
        original_filename = os.path.basename(report.file_path)
        if original_filename.startswith(f"{user.id}_"):
            original_filename = original_filename[len(f"{user.id}_"):]

        return open(report.file_path, "rb"), original_filename, content_type

    async def get_file_metadata(
        self,
        user: User,
        report_id: int
    ) -> dict:
        """Get file metadata.
        
        Returns:
            dict: File metadata including size, type, and last modified time
        """
        report = (
            self.db.query(Report)
            .filter(Report.id == report_id, Report.user_id == user.id)
            .first()
        )
        if not report:
            raise HTTPException(status_code=404, detail="Report not found")

        if not os.path.exists(report.file_path):
            raise HTTPException(status_code=404, detail="File not found")

        file_stat = os.stat(report.file_path)
        return {
            "size": file_stat.st_size,
            "type": report.file_type,
            "last_modified": file_stat.st_mtime,
            "created_at": file_stat.st_ctime,
            "path": report.file_path
        }

    async def stream_file(
        self,
        user: User,
        report_id: int
    ) -> StreamingResponse:
        """Stream file content.
        
        Returns:
            StreamingResponse: FastAPI streaming response
        """
        file_obj, filename, content_type = await self.get_file_content(user, report_id)
        
        return StreamingResponse(
            file_obj,
            media_type=content_type,
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"'
            }
        )

    async def _handle_file_upload(self, report_id: uuid.UUID, file: UploadFile) -> None:
        """Handle file upload for a report."""
        try:
            # Validate file size
            file_size = 0
            chunk_size = 1024 * 1024  # 1MB chunks
            while chunk := await file.read(chunk_size):
                file_size += len(chunk)
                if file_size > self.max_upload_size:
                    raise ValidationException(f"File size exceeds maximum allowed size of {self.max_upload_size} bytes")
            
            # Validate file extension
            file_extension = os.path.splitext(file.filename)[1].lower().lstrip('.')
            if file_extension not in self.allowed_extensions:
                raise ValidationException(f"File extension not allowed. Allowed extensions: {', '.join(self.allowed_extensions)}")
            
            # Create upload directory if it doesn't exist
            try:
                os.makedirs(self.upload_dir, exist_ok=True)
            except OSError as e:
                raise DatabaseError(f"Error creating upload directory: {str(e)}")
            
            # Save file
            file_path = os.path.join(self.upload_dir, f"{report_id}_{file.filename}")
            try:
                with open(file_path, "wb") as buffer:
                    await file.seek(0)
                    while chunk := await file.read(chunk_size):
                        buffer.write(chunk)
            except IOError as e:
                raise DatabaseError(f"Error saving file: {str(e)}")
        except ValidationException as e:
            raise e
        except DatabaseError as e:
            raise e
        except Exception as e:
            logger.error(f"Unexpected error handling file upload: {str(e)}", exc_info=True)
            raise DatabaseError(f"Error handling file upload: {str(e)}")

    async def _delete_report_files(self, report_id: uuid.UUID) -> None:
        """Delete files associated with a report."""
        try:
            # Find all files for this report
            for filename in os.listdir(self.upload_dir):
                if filename.startswith(str(report_id)):
                    file_path = os.path.join(self.upload_dir, filename)
                    if os.path.isfile(file_path):
                        os.remove(file_path)
        except Exception as e:
            logger.error(f"Error deleting report files: {str(e)}")

    async def process_report(self, db: Session, report_id: uuid.UUID) -> None:
        """Process a report asynchronously."""
        report = None
        try:
            report = self.db.query(Report).filter(Report.id == report_id).first()
            if not report:
                raise NotFoundException("Report not found")

            # Update status to processing
            report.status = "processing"
            self.db.add(report)
            self.db.commit()

            # Process based on file type
            try:
                content = await self._process_file(report)
            except Exception as e:
                raise AIProcessingError(f"Error processing file content: {str(e)}")

            # Update report with processed content
            update_data = {
                "content": content,
                "status": "processed",
                "processed_at": datetime.utcnow()
            }
            self.db.add(report)
            self.db.commit()

        except NotFoundException as e:
            raise e
        except AIProcessingError as e:
            if report:
                report.status = "failed"
                report.error = str(e)
                self.db.add(report)
                self.db.commit()
            raise e
        except Exception as e:
            logger.error(f"Unexpected error processing report: {str(e)}", exc_info=True)
            if report:
                report.status = "failed"
                report.error = str(e)
                self.db.add(report)
                self.db.commit()
            raise DatabaseError(f"Error processing report: {str(e)}")

    async def _process_file(self, report: Report) -> Dict[str, Any]:
        """Process the report file based on its type."""
        if not os.path.exists(report.file_path):
            raise FileNotFoundError(f"Report file not found: {report.file_path}")

        try:
            if report.file_type == "pdf":
                return await self._process_pdf(report)
            elif report.file_type in ["docx", "doc"]:
                return await self._process_docx(report)
            elif report.file_type in ["xlsx", "xls"]:
                return await self._process_excel(report)
            elif report.file_type == "csv":
                return await self._process_csv(report)
            else:
                raise ValueError(f"Unsupported file type: {report.file_type}")
        except Exception as e:
            raise ValidationException(f"Error processing file: {str(e)}")

    async def _process_pdf(self, report: Report) -> Dict[str, Any]:
        """Process PDF files."""
        # TODO: Implement PDF processing
        raise NotImplementedError("PDF processing not implemented")

    async def _process_docx(self, report: Report) -> Dict[str, Any]:
        """Process DOCX files."""
        doc = Document(report.file_path)
        return {
            "text": "\n".join([p.text for p in doc.paragraphs if p.text.strip()]),
            "tables": [table.text for table in doc.tables],
            "metadata": {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections)
            }
        }

    async def _process_excel(self, report: Report) -> Dict[str, Any]:
        """Process Excel files."""
        wb = openpyxl.load_workbook(report.file_path)
        data = {}
        for sheet in wb.worksheets:
            sheet_data = []
            for row in sheet.iter_rows(values_only=True):
                sheet_data.append([str(cell) for cell in row if cell is not None])
            data[sheet.title] = sheet_data
        return {
            "sheets": data,
            "metadata": {
                "sheets": len(wb.worksheets),
                "active_sheet": wb.active.title
            }
        }

    async def _process_csv(self, report: Report) -> Dict[str, Any]:
        """Process CSV files."""
        df = pd.read_csv(report.file_path)
        return {
            "data": df.to_dict(orient="records"),
            "metadata": {
                "rows": len(df),
                "columns": list(df.columns),
                "dtypes": df.dtypes.to_dict()
            }
        }

    async def create_version(
        self,
        db: Session,
        report: Report,
        changes_description: str
    ) -> ReportVersion:
        """Create a new version of the report."""
        # Get the latest version number
        latest_version = (
            db.query(ReportVersion)
            .filter(ReportVersion.report_id == report.id)
            .order_by(ReportVersion.version_number.desc())
            .first()
        )
        version_number = (latest_version.version_number + 1) if latest_version else 1

        # Create new version
        version = ReportVersion(
            report_id=report.id,
            version_number=version_number,
            file_path=report.file_path,
            changes_description=changes_description
        )
        db.add(version)
        db.commit()
        db.refresh(version)

        return version

    async def update_metadata(
        self,
        db: Session,
        report: Report,
        metadata: Dict[str, Any]
    ) -> None:
        """Update report metadata."""
        report.metadata.update(metadata)
        report.metadata["last_updated"] = datetime.utcnow().isoformat()
        db.add(report)
        db.commit()
        db.refresh(report)

# Create service instance
report_service = ReportService(None) 